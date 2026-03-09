"""
Job file artifact generator — produces downloadable Word/Excel from job output.
Used when job pipeline completes (JOB_READY) to offer file download in JobSplitView.
"""
import logging
import os
import re
from datetime import datetime
from typing import Any, Dict, List, Literal, Optional

logger = logging.getLogger(__name__)

OUTPUT_DIR = "/tmp/job_artifacts"


def parse_output_to_sections(text: str) -> List[Dict[str, str]]:
    """
    Parse agent output (markdown with ## / ### headers) into sections for Word.
    Fallback: single section "Inhoud" with full text if no headers.
    """
    if not text or not isinstance(text, str):
        return []
    sections: List[Dict[str, str]] = []
    current_heading = "Inhoud"
    current_body: List[str] = []

    for line in text.split("\n"):
        if line.startswith("## ") or line.startswith("### "):
            if current_body:
                sections.append({
                    "heading": current_heading,
                    "body": "\n".join(current_body).strip(),
                })
            current_heading = line.lstrip("#").strip()
            current_body = []
        else:
            current_body.append(line)

    if current_body:
        sections.append({
            "heading": current_heading,
            "body": "\n".join(current_body).strip(),
        })

    if not sections:
        sections.append({"heading": "Inhoud", "body": text.strip()})

    return sections


def _generate_docx(content: Dict[str, Any], filepath: str) -> None:
    """Generate Word document from content dict."""
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    title = doc.add_heading(content.get("title", "Rapport"), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run(f"{content.get('brand', '')} | {content.get('date', '')}").italic = True
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    for section in content.get("sections", []):
        doc.add_heading(section.get("heading", ""), level=1)
        doc.add_paragraph(section.get("body", ""))

    doc.save(filepath)


def _generate_xlsx(content: Dict[str, Any], filepath: str) -> None:
    """Generate Excel from content dict."""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = openpyxl.Workbook()
    if wb.active:
        wb.remove(wb.active)

    for sheet_data in content.get("sheets", []):
        ws = wb.create_sheet(sheet_data.get("name", "Sheet"))

        headers = sheet_data.get("headers", [])
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True, color="FFFFFF", name="Arial")
            cell.fill = PatternFill(start_color="2D2D2D", end_color="2D2D2D", fill_type="solid")
            cell.alignment = Alignment(horizontal="center")

        for row_idx, row in enumerate(sheet_data.get("rows", []), 2):
            for col_idx, value in enumerate(row, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)

        if headers:
            ws.auto_filter.ref = ws.dimensions
        ws.freeze_panes = "A2"

    wb.save(filepath)


async def generate_job_artifact(
    conn,
    job_id: str,
    content: Dict[str, Any],
    file_type: Literal["docx", "xlsx"],
    brand_name: str,
    job_title: str,
) -> Dict[str, str]:
    """
    Generate Word or Excel file from job output. Saves path in jobs table.
    Uses jobs.id (not job_id column) — job_id is the UUID value for the id column.
    Returns: {"path": str, "filename": str, "type": str}
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    timestamp = datetime.now().strftime("%Y-%m-%d")
    safe_brand = "".join(c if c.isalnum() or c in " -_" else "_" for c in brand_name)[:30]
    safe_title = (job_title or "Output")[:30].replace(" ", "_")
    filename = f"{safe_title}_{safe_brand}_{timestamp}.{file_type}"
    filepath = os.path.join(OUTPUT_DIR, filename)

    if file_type == "docx":
        _generate_docx(content, filepath)
    elif file_type == "xlsx":
        _generate_xlsx(content, filepath)
    else:
        raise ValueError(f"Unsupported file_type: {file_type}")

    # ASSUMPTION: jobs table uses id as PK; job_id is the UUID value for id
    await conn.execute(
        """
        UPDATE jobs
        SET file_artifact_path = $1, file_artifact_type = $2, file_artifact_name = $3
        WHERE id = $4
        """,
        filepath,
        file_type,
        filename,
        job_id,
    )

    logger.info("Generated job artifact: %s", filepath)
    return {"path": filepath, "filename": filename, "type": file_type}
